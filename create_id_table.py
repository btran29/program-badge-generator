from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import csv


def import_csv(csv_filename):
    """
    Imports a two-column tab-delimited text file as a variable. Skips the header / 1st row.

    Should be structured as such:

    Name    Affiliation
    Name1   Affiliation1
    Name2   Affiliation2
    ..
    NameX   AffiliationX
    etc.

    File format: txt, tab-delimited, UTF-8

    :param csv_filename:
    :type csv_filename: str
    :return:
    """
    records = []
    # Open CSV, iterate making IDs through each row
    with open(csv_filename, 'rb') as csv_file:
        name_reader = csv.reader(csv_file, delimiter='\t', quotechar='|')
        for num, row in enumerate(name_reader, start=0):
            if num >= 1:
                current_tuple = (row[0], row[1])
                records.append(current_tuple)
    records = tuple(records)
    return records


def make_id_document(document_object, record, event, style):
    """
    Main function to make that arranges ID cards into a word document.

    Currently set to 6 IDs per page.

    Required libraries: docx

    :param document_object: template document
    :param record: Imported records from import function, tuple of tuples
    :param event: Name of event header to be on every ID
    :param style: Table style present in template document
    :type record: tuple
    :type event: str
    :type style: str
    :return:
    """
    id_counter = 0
    for name, affiliation in record:

        # Make ID
        make_id_table(event, name, affiliation, document_object, style)
        docx_line_break(document_object)

        # Counter
        id_counter += 1

        # Page break every 6 IDs
        if id_counter % 6 == 0:
            docx_page_break(document_object)
            docx_line_break(document_object)
            continue

        # Column break every 3 IDs
        if id_counter % 3 == 0:
            docx_column_break(document_object)

    # Fix size
    fix_id_table_size(document_object)


def docx_line_break(document_object):
    p = document_object.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.LINE)


def docx_column_break(document_object):
    p = document_object.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.COLUMN)


def docx_page_break(document_object):
    p = document_object.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def make_id_table(event, name, affiliation, document_object, style):
    """
    Make an individual ID given specific strings.
    :param event:
    :param name:
    :param affiliation:
    :param document_object:
    :param style
    :type event: str
    :type name: str
    :type affiliation: str
    :type document_object: python docx document object
    :type style: table style existing in template document
    :return:
    """
    id_table = document_object.add_table(3, 1, style=style)
    id_table.cell(0, 0).text = event
    id_table.cell(1, 0).text = str(name)
    id_table.cell(2, 0).text = str(affiliation)


def fix_id_table_size(document_object):
    """
    Adjust Height / column for each table

    Required classes: docx.shared.Inches
    :param document_object
    :type document_object: python docx document object
    :return:
    """
    for table in document_object.tables:
        for num, row in enumerate(table.rows, start=0):
            if num is 0:
                row.height = Inches(.5)
            if num is 1:
                row.height = Inches(1)
            if num is 2:
                row.height = Inches(1)
        for column in table.columns:
            column.height = Inches(4)


###
#
# USAGE
#
###
#
#
# # Import relevant libraries
#
#
# # Set event and style strings
# test_event = "2019 UCI MSTP RETREAT"
# test_style = "Badge"
#
# # Open template document
# test_document = Document("template_test.docx")
#
# # Import names list
# records = import_csv("names_list_FINAL.txt")
# 
# # Run script
# make_id_document(test_document, records, test_event, test_style)
#
# # Save
# test_document.save("template_test1.docx")
